import re
from docx import Document


class DocxProjectParser:
    """Парсер реальных ТЗ ФМБА (Word-документы) с разметкой достоверности и трассировкой источников."""

    DATE_PATTERN = re.compile(r'\b\d{2}\.\d{2}\.\d{4}\b')

    def __init__(self, file_path_or_buffer):
        self.doc = Document(file_path_or_buffer)

    def parse(self) -> dict:
        full_text = self._extract_full_text_with_sources()
        flat_text = "\n".join([item['text'] for item in full_text])

        # Блок А: Реквизиты проекта
        project_fields = self._parse_requisites(flat_text)

        # Блок Б и В: Этапы, задачи и выходы из таблицы (пункт 11)
        stages = self._parse_stages_and_outputs()

        return {
            "project_fields": project_fields,
            "stages": stages
        }

    def _extract_full_text_with_sources(self) -> list[dict]:
        """Собирает текст абзацев с указанием источника (для трассировки)."""
        content = []
        for i, p in enumerate(self.doc.paragraphs):
            txt = p.text.strip()
            if txt:
                content.append({
                    "text": txt,
                    "source": f"Абзац {i+1}"
                })
        return content

    def _parse_requisites(self, text: str) -> dict:
        """Извлечение основных реквизитов проекта по ключевым словам ТЗ ФМБА."""
        
        # 1. Наименование проекта
        name_match = re.search(r'по теме:\s*«([^»]+)»', text, re.IGNORECASE)
        шифр_match = re.search(r'шифр:\s*«([^»]+)»', text, re.IGNORECASE)
        
        project_name = name_match.group(1).strip() if name_match else ""
        if шифр_match:
            project_name += f" (шифр: {шифр_match.group(1).strip()})"

        # 2. Сроки выполнения (Пункт 6)
        dates = self.DATE_PATTERN.findall(text)
        start_date = dates[0] if len(dates) > 0 else None
        end_date = dates[1] if len(dates) > 1 else None

        # 3. Заказчик (Пункт 4)
        customer_match = re.search(r'4\.\s*Государственный заказчик:\s*([^\n]+)', text, re.IGNORECASE)
        customer = customer_match.group(1).strip() if customer_match else "Федеральное медико-биологическое агентство"

        # 4. Исполнитель (Пункт 5)
        executor_match = re.search(r'5\.\s*Исполнитель[^:]*:\s*([^\n]+)', text, re.IGNORECASE)
        executor = executor_match.group(1).strip() if executor_match else ""

        return {
            "name": {
                "value": project_name or "Не найдено",
                "status": "Распознано" if project_name else "Не найдено",
                "source": "Преамбула / Тема ТЗ"
            },
            "customer": {
                "value": customer,
                "status": "Распознано" if customer else "Не найдено",
                "source": "Пункт 4 ТЗ"
            },
            "executor": {
                "value": executor,
                "status": "Распознано" if executor else "Требует проверки",
                "source": "Пункт 5 ТЗ"
            },
            "start_date": {
                "value": start_date,
                "status": "Распознано" if start_date else "Не найдено",
                "source": "Пункт 6 ТЗ (Сроки выполнения)"
            },
            "end_date": {
                "value": end_date,
                "status": "Распознано" if end_date else "Не найдено",
                "source": "Пункт 6 ТЗ (Сроки выполнения)"
            }
        }

    def _parse_stages_and_outputs(self) -> list[dict]:
        """Парсинг календарного плана (Таблица этапов, обычно пункт 11)."""
        stages_map = {}
        
        for table_idx, table in enumerate(self.doc.tables):
            if len(table.rows) <= 1:
                continue

            # Проверяем, похожа ли таблица на календарный план по заголовкам
            header_text = " ".join([cell.text.strip().lower() for cell in table.rows[0].cells])
            if not ("этап" in header_text or "наименование" in header_text or "срок" in header_text):
                continue

            current_main_stage = None

            for row_idx in range(1, len(table.rows)):
                cells = [c.text.strip() for c in table.rows[row_idx].cells]
                if len(cells) < 4:
                    continue

                num_col = cells[0]
                name_col = cells[1]
                output_col = cells[2]
                dates_col = cells[3]

                if not num_col and not name_col:
                    continue

                # Извлекаем даты из строки таблицы
                row_dates = self.DATE_PATTERN.findall(dates_col)
                start_d = row_dates[0] if len(row_dates) > 0 else None
                end_d = row_dates[1] if len(row_dates) > 1 else start_d

                # Определяем, главный это этап (например, "1") или подэтап ("1.1")
                is_substage = "." in num_col

                if not is_substage:
                    # Главный этап
                    stage_key = num_col
                    stages_map[stage_key] = {
                        "name": f"Этап {num_col}. {name_col}",
                        "start_date": start_d,
                        "end_date": end_d,
                        "status": "Распознано" if start_d and end_d else "Распознано с предупреждением",
                        "source": f"Таблица плана, строка {row_idx + 1}",
                        "tasks": []
                    }
                    current_main_stage = stage_key
                else:
                    # Подэтап трактуем как Задачу внутри текущего главного этапа
                    task_data = {
                        "name": f"Задача {num_col}. {name_col}",
                        "start_date": start_d,
                        "end_date": end_d,
                        "status": "Распознано" if start_d and end_d else "Распознано с предупреждением",
                        "source": f"Таблица плана, строка {row_idx + 1}",
                        "outputs": []
                    }

                    # Парсим результаты (выходы), если они указаны в 3 колонке
                    if output_col and output_col != "-":
                        output_items = [o.strip() for o in output_col.split('\n') if o.strip()]
                        for item in output_items:
                            # Определяем тип выхода по ключевым словам
                            out_type = "Отчет"
                            low_item = item.lower()
                            if "стать" in low_item:
                                out_type = "Статья"
                            elif "патент" in low_item or "рид" in low_item or "ноу-хау" in low_item:
                                out_type = "Патент/РИД"
                            elif "акт" in low_item or "справк" in low_item:
                                out_type = "Акт/Справка"

                            task_data["outputs"].append({
                                "name": item,
                                "output_type": out_type,
                                "deadline": end_d,
                                "status": "Распознано"
                            })

                    # Привязываем задачу к этапу
                    parent_key = num_col.split('.')[0]
                    if parent_key in stages_map:
                        stages_map[parent_key]["tasks"].append(task_data)
                    elif current_main_stage and current_main_stage in stages_map:
                        stages_map[current_main_stage]["tasks"].append(task_data)

        return list(stages_map.values())