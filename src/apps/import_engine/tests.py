import io
from django.test import TestCase
import pandas as pd
from django.contrib.auth import get_user_model
from docx import Document
from apps.import_engine.parser_docx import DocxProjectParser
from apps.import_engine.services import import_staff_structure
from apps.organization.models import DepartmentMembership

User = get_user_model()


class DocxParserTestCase(TestCase):
    """Тест-кейс для проверки извлечения реквизитов и структуры этапов из .docx файлов."""

    def _create_sample_docx(self) -> io.BytesIO:
        """Создает в памяти мини-копию ТЗ для тестирования парсера."""
        doc = Document()
        
        # Реквизиты
        doc.add_paragraph("ТЕХНИЧЕСКОЕ ЗАДАНИЕ")
        doc.add_paragraph("на выполнение прикладной научно-исследовательской работы")
        doc.add_paragraph("по теме: «Создание системы управления биобанкингом» (шифр: «Депозит 2025»)")
        doc.add_paragraph("4. Государственный заказчик: Федеральное медико-биологическое агентство")
        doc.add_paragraph("5. Исполнитель: ФГБУН ЮУрФНКЦ МБ ФМБА России")
        doc.add_paragraph("6. Сроки выполнения работы: начало - 01.04.2025; окончание - 31.12.2027.")

        # Таблица календарного плана (Пункт 11)
        table = doc.add_table(rows=3, cols=4)
        
        # Шапка таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Номер этапа"
        hdr_cells[1].text = "Наименование"
        hdr_cells[2].text = "Результат"
        hdr_cells[3].text = "Сроки"

        # Главный этап (1)
        row1 = table.rows[1].cells
        row1[0].text = "1"
        row1[1].text = "Разработка системы управления банком образцов."
        row1[2].text = "-"
        row1[3].text = "01.04.2025 - 31.12.2025"

        # Подэтап / Задача (1.1)
        row2 = table.rows[2].cells
        row2[0].text = "1.1"
        row2[1].text = "Разработка стратегии поиска и сбора информации."
        row2[2].text = "Аналитический обзор о предикторах.\nПротокол стратегии, утвержденный на Ученом совете."
        row2[3].text = "01.04.2025 - 30.06.2025"

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def test_docx_parser_extracts_requisites_and_stages(self):
        docx_buffer = self._create_sample_docx()
        parser = DocxProjectParser(docx_buffer)
        result = parser.parse()

        # 1. Проверяем наличие основных блоков словаря
        self.assertIn("project_fields", result)
        self.assertIn("stages", result)

        # 2. Проверяем корректность извлечения реквизитов
        fields = result["project_fields"]
        self.assertEqual(fields["start_date"]["value"], "01.04.2025")
        self.assertEqual(fields["end_date"]["value"], "31.12.2027")
        self.assertIn("Создание системы управления биобанкингом", fields["name"]["value"])
        self.assertEqual(fields["customer"]["value"], "Федеральное медико-биологическое агентство")

        # 3. Проверяем иерархию этапов и задач
        stages = result["stages"]
        self.assertTrue(len(stages) > 0, "Парсер должен найти хотя бы один главный этап")
        
        first_stage = stages[0]
        self.assertIn("Этап 1", first_stage["name"])
        self.assertEqual(first_stage["start_date"], "01.04.2025")
        self.assertEqual(first_stage["end_date"], "31.12.2025")

        # Проверяем наличие вложенных задач (подэтапов)
        self.assertTrue(len(first_stage["tasks"]) > 0, "Главный этап должен содержать задачи/подэтапы")
        first_task = first_stage["tasks"][0]
        self.assertIn("Задача 1.1", first_task["name"])
        self.assertEqual(first_task["start_date"], "01.04.2025")
        self.assertEqual(first_task["end_date"], "30.06.2025")

        # Проверяем извлечение выходных результатов (выходов)
        self.assertTrue(len(first_task["outputs"]) > 0, "У задачи должны распознаться выходы из 3-й колонки")
        output_names = [out["name"] for out in first_task["outputs"]]
        self.assertTrue(any("Аналитический обзор" in name for name in output_names))

class StaffImportServiceTestCase(TestCase):
    """Тест-кейс для проверки импорта штатного расписания и логики совмещения."""

    def _create_sample_excel(self, rows_data) -> io.BytesIO:
        df = pd.DataFrame(rows_data)
        excel_buffer = io.BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Sheet1')
        excel_buffer.seek(0)
        return excel_buffer

    def test_import_staff_primary_conflict_resolution(self):
            data = [
                {
                    "ФИО сотрудника": "Иванов Иван Иванович",
                    "Электронная почта / Email": "ivanov@inst.ru",
                    "Логин / Username": "ivanov",
                    "Наименование подразделения": "НИИ ЯФ / Лаборатория физики",
                    "Наименование должности": "Младший научный сотрудник",
                    "Признак основного места работы": "Нет"
                },
                {
                    "ФИО сотрудника": "Иванов Иван Иванович",
                    "Электронная почта / Email": "ivanov@inst.ru",
                    "Логин / Username": "ivanov",
                    "Наименование подразделения": "НИИ ЯФ / Отдел ИТ",
                    "Наименование должности": "Старший научный сотрудник",
                    "Признак основного места работы": "Да"
                }
            ]

            excel_file = self._create_sample_excel(data)
            result = import_staff_structure(actor=None, file=excel_file)

            self.assertEqual(result["status"], "success")
            self.assertEqual(result["processed_rows"], 2)
            self.assertEqual(result["created_users_count"], 1)

            # ДОБАВЛЕН ИМПОРТ И ПОИСК EMPLOYEE
            from apps.organization.models import Employee
            
            user = User.objects.get(email="ivanov@inst.ru")
            employee = Employee.objects.get(user=user)
            
            # ИСПРАВЛЕНА ФИЛЬТРАЦИЯ С user=user НА employee=employee
            memberships = DepartmentMembership.objects.filter(employee=employee)
            
            self.assertEqual(memberships.count(), 2)

            primary_count = memberships.filter(is_primary=True).count()
            self.assertEqual(primary_count, 1, "У сотрудника должно быть ровно одно основное место работы")
            
            it_membership = memberships.get(department__name="Отдел ИТ")
            self.assertTrue(it_membership.is_primary)